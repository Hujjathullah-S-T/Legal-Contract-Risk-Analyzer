import html
import io
import re
from collections import Counter
from datetime import datetime
from difflib import SequenceMatcher

from PyPDF2 import PdfReader
from docx import Document
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

RISK_LIBRARY = {
    "High Risk": {
        "weight": 24,
        "keywords": {
            "penalty": "Penalty language may impose immediate financial exposure.",
            "breach": "Breach terminology can trigger enforcement or damages.",
            "liability": "Liability wording may expand legal or financial responsibility.",
            "terminate immediately": "Immediate termination can create abrupt operational risk.",
            "unlimited liability": "Unlimited liability can expose a party to uncapped losses.",
            "damages": "Damages clauses can increase financial exposure after disputes.",
            "default": "Default wording can accelerate remedies against a party.",
            "exclusive jurisdiction": "Exclusive jurisdiction can reduce negotiation flexibility in disputes.",
        },
    },
    "Medium Risk": {
        "weight": 14,
        "keywords": {
            "delay": "Delay obligations can create service-level or delivery exposure.",
            "obligation": "Obligation wording may create enforceable performance duties.",
            "indemnify": "Indemnity language can shift third-party risk.",
            "notice period": "Notice periods affect termination flexibility and operational timing.",
            "warranty": "Warranty promises may create repair, refund, or compliance exposure.",
            "compliance": "Compliance obligations may require additional controls or audits.",
        },
    },
    "Low Risk": {
        "weight": 6,
        "keywords": {
            "may": "Optional wording may introduce ambiguity in responsibilities.",
            "should": "Soft language can weaken enforceability or expectations.",
            "optional": "Optional terms may reduce clarity around commitments.",
            "reasonable efforts": "Reasonable efforts language can be subjective in disputes.",
        },
    },
}

CLAUSE_LIBRARY = {
    "Payment Clause": ["pay", "payment", "invoice", "fee", "fees", "amount", "consideration"],
    "Termination Clause": ["terminate", "termination", "notice period", "expire", "expiration"],
    "Liability Clause": ["liability", "damages", "indemnify", "indemnity", "hold harmless"],
    "Confidentiality Clause": ["confidential", "non-disclosure", "nda", "proprietary information", "trade secret"],
    "Dispute Resolution Clause": ["arbitration", "dispute", "jurisdiction", "governing law", "court"],
    "Data Privacy Clause": ["personal data", "gdpr", "privacy", "consent", "data processing"],
}

MISSING_CLAUSE_TARGETS = ["Termination Clause", "Dispute Resolution Clause"]

AMBIGUOUS_TERMS = {
    "reasonable": "Reasonable is open to interpretation and may create ambiguity.",
    "as soon as possible": "This phrase has no objective deadline.",
    "best effort": "Best effort is subjective and hard to enforce consistently.",
    "best efforts": "Best efforts is subjective and hard to enforce consistently.",
    "from time to time": "This phrase can make obligations unclear over time.",
    "promptly": "Promptly may be too vague without a measurable deadline.",
}

RISK_RECOMMENDATIONS = {
    "unlimited liability": "Replace with a capped liability clause tied to contract value.",
    "liability": "Consider limiting liability to direct damages with a monetary cap.",
    "indemnify": "Clarify indemnity scope, triggers, and exclusions.",
    "terminate immediately": "Add cure periods before immediate termination applies.",
    "penalty": "Use a proportionate liquidated damages clause instead of a broad penalty.",
    "exclusive jurisdiction": "Consider neutral jurisdiction or arbitration if flexibility is needed.",
}

COMPLIANCE_RULES = [
    {
        "name": "Privacy notice gap",
        "trigger": ["personal data", "customer data", "sensitive data"],
        "required_any": ["consent", "privacy", "gdpr", "security", "confidential"],
        "message": "Personal data is referenced without a strong privacy/compliance safeguard.",
    },
    {
        "name": "Security safeguard gap",
        "trigger": ["data", "system", "access"],
        "required_any": ["encryption", "security", "access control", "confidential"],
        "message": "The contract references data or system access but does not clearly mention security controls.",
    },
]

QUESTION_TOPICS = {
    "penalty": ["penalty", "liquidated damages"],
    "liability": ["liability", "damages", "indemnify", "indemnity"],
    "termination": ["terminate", "termination", "notice period"],
    "payment": ["pay", "payment", "invoice", "fee", "amount"],
    "confidentiality": ["confidential", "non-disclosure", "trade secret"],
    "privacy": ["privacy", "personal data", "gdpr", "consent"],
}

ROLE_DESCRIPTIONS = {
    "Lawyer": "Detailed legal signals, clause intelligence, and sentence-level reasoning.",
    "Client": "Plain-language explanation with practical concerns and missing safeguards.",
    "Manager": "Fast risk summary with score, top issues, and actions.",
}


def contains_term(text, term):
    return bool(re.search(rf"\b{re.escape(term)}\b", text, re.IGNORECASE))


def normalize_text(text):
    return re.sub(r"\s+", " ", text or "").strip()


def split_sentences(text):
    cleaned = normalize_text(text)
    if not cleaned:
        return []
    return [item.strip() for item in re.split(r"(?<=[.!?])\s+", cleaned) if item.strip()]


def level_priority(level):
    return {"High Risk": 0, "Medium Risk": 1, "Low Risk": 2}[level]


def level_class(level):
    return level.lower().split()[0]


def read_uploaded_file(uploaded_file):
    file_name = uploaded_file.name.lower()
    if file_name.endswith(".txt"):
        return uploaded_file.getvalue().decode("utf-8", errors="ignore")
    if file_name.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(uploaded_file.getvalue()))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if file_name.endswith(".docx"):
        document = Document(io.BytesIO(uploaded_file.getvalue()))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
    return ""


def detect_clause_types(text):
    detected = []
    for clause_name, patterns in CLAUSE_LIBRARY.items():
        matches = [pattern for pattern in patterns if contains_term(text, pattern)]
        if matches:
            detected.append(
                {
                    "clause": clause_name,
                    "matches": matches,
                    "explanation": f"Detected through terms like {', '.join(matches[:3])}.",
                }
            )
    return detected


def build_clause_cards(clauses, findings, text):
    cards = []
    clause_lookup = {clause["clause"]: clause for clause in clauses}
    lower_text = text.lower()
    for clause_name in CLAUSE_LIBRARY:
        clause = clause_lookup.get(clause_name)
        clause_matches = clause["matches"] if clause else []
        relevant_findings = []
        for finding in findings:
            if finding["term"] in lower_text and (
                any(pattern in finding["term"] or finding["term"] in pattern for pattern in CLAUSE_LIBRARY[clause_name])
                or clause_name in {"Liability Clause", "Termination Clause", "Payment Clause"}
                and clause_matches
            ):
                relevant_findings.append(finding)
        if relevant_findings:
            primary = sorted(relevant_findings, key=lambda item: level_priority(item["level"]))[0]
            risk_level = primary["level"]
            explanation = primary["explanation"]
            recommendation = RISK_RECOMMENDATIONS.get(
                primary["term"],
                f"Review the {clause_name.lower()} and tighten the wording around {', '.join(clause_matches[:2]) or 'detected obligations'}.",
            )
            score = min(sum(item["weight"] for item in relevant_findings), 100)
        elif clause:
            risk_level = "Low Risk"
            explanation = clause["explanation"]
            recommendation = f"Keep the {clause_name.lower()} precise and aligned with negotiated obligations."
            score = 12
        else:
            risk_level = "Needs Review"
            explanation = f"{clause_name} was not clearly detected in this contract."
            recommendation = f"Consider adding a clear {clause_name.lower()} to improve legal coverage."
            score = 0
        cards.append(
            {
                "clause": clause_name,
                "risk_level": risk_level,
                "score": score,
                "matches": clause_matches,
                "explanation": explanation,
                "recommendation": recommendation,
            }
        )
    return cards


def extract_entities(text):
    parties = set()
    dates = set()
    money = set()

    party_patterns = [
        r"\b(?:Client|Company|Vendor|Supplier|Consultant|Employee|Contractor|Buyer|Seller|Lessor|Lessee|Service Provider|Customer)\b",
        r"\b[A-Z][A-Za-z&., ]+(?:Ltd|Limited|LLP|LLC|Inc|Corp|Corporation|Private Limited|Pvt\. Ltd\.)\b",
    ]
    date_patterns = [
        r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b",
        r"\b\d{1,2}(?:st|nd|rd|th)?\s+(?:Jan|January|Feb|February|Mar|March|Apr|April|May|Jun|June|Jul|July|Aug|August|Sep|Sept|September|Oct|October|Nov|November|Dec|December)\b",
        r"\b(?:Jan|January|Feb|February|Mar|March|Apr|April|May|Jun|June|Jul|July|Aug|August|Sep|Sept|September|Oct|October|Nov|November|Dec|December)\s+\d{1,2},?\s+\d{4}\b",
    ]
    money_patterns = [
        r"(?:₹|Rs\.?\s?|INR\s?)\d[\d,]*(?:\.\d+)?",
        r"(?:\$|USD\s?)\d[\d,]*(?:\.\d+)?",
        r"(?:€|EUR\s?)\d[\d,]*(?:\.\d+)?",
    ]

    for pattern in party_patterns:
        parties.update(re.findall(pattern, text))
    for pattern in date_patterns:
        dates.update(re.findall(pattern, text, flags=re.IGNORECASE))
    for pattern in money_patterns:
        money.update(re.findall(pattern, text, flags=re.IGNORECASE))

    return {
        "Parties": sorted({item.strip() for item in parties if item.strip()}),
        "Dates": sorted({item.strip() for item in dates if item.strip()}),
        "Money Values": sorted({item.strip() for item in money if item.strip()}),
    }


def analyze_text(text):
    findings = []
    for level, config in RISK_LIBRARY.items():
        for term, explanation in config["keywords"].items():
            if contains_term(text, term):
                findings.append(
                    {
                        "term": term,
                        "level": level,
                        "weight": config["weight"],
                        "explanation": explanation,
                    }
                )
    return findings


def analyze_sentences(text):
    sentence_findings = []
    for sentence in split_sentences(text):
        matches = []
        sentence_score = 0
        for level, config in RISK_LIBRARY.items():
            for term, explanation in config["keywords"].items():
                if contains_term(sentence, term):
                    matches.append({"term": term, "level": level, "reason": explanation})
                    sentence_score += config["weight"]
        if matches:
            primary_level = sorted(matches, key=lambda item: level_priority(item["level"]))[0]["level"]
            sentence_findings.append(
                {
                    "sentence": sentence,
                    "matches": matches,
                    "score": min(sentence_score, 100),
                    "level": primary_level,
                    "reason": " ".join(match["reason"] for match in matches[:2]),
                }
            )
    return sentence_findings


def detect_ambiguity(text):
    findings = []
    for sentence in split_sentences(text):
        for term, message in AMBIGUOUS_TERMS.items():
            if contains_term(sentence, term):
                findings.append({"term": term, "sentence": sentence, "message": message})
    return findings


def extract_obligations(text):
    obligations = []
    pattern = re.compile(
        r"\b([A-Z][A-Za-z ]{1,40}|client|vendor|company|supplier|buyer|seller|employee|contractor|consultant)\b\s+"
        r"(shall|must|will|agrees to|is required to)\s+"
        r"([a-z]+)\s+([^.!?]+)",
        flags=re.IGNORECASE,
    )
    for sentence in split_sentences(text):
        match = pattern.search(sentence)
        if match:
            obligations.append(
                {
                    "subject": match.group(1).strip().title(),
                    "action": match.group(3).strip().capitalize(),
                    "object": match.group(4).strip(),
                    "sentence": sentence,
                }
            )
    return obligations


def detect_clause_dependencies(text):
    dependencies = []
    for sentence in split_sentences(text):
        lowered = sentence.lower()
        if ("terminate" in lowered or "termination" in lowered) and ("pay" in lowered or "payment" in lowered):
            if any(trigger in lowered for trigger in ["if", "unless", "upon", "only if", "fails", "failure"]):
                dependencies.append(
                    {
                        "source": "Termination Clause",
                        "depends_on": "Payment Clause",
                        "sentence": sentence,
                        "reason": "Termination appears conditional on payment behavior.",
                    }
                )
        if "liability" in lowered and ("confidential" in lowered or "data" in lowered):
            dependencies.append(
                {
                    "source": "Liability Clause",
                    "depends_on": "Confidentiality/Data Privacy Clause",
                    "sentence": sentence,
                    "reason": "Liability appears linked to confidentiality or data handling obligations.",
                }
            )
    return dependencies


def detect_sensitive_data(text):
    findings = []
    patterns = {
        "Email": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b",
        "Phone": r"\b(?:\+91[- ]?)?[6-9]\d{9}\b",
        "PAN": r"\b[A-Z]{5}\d{4}[A-Z]\b",
    }
    for label, pattern in patterns.items():
        for match in re.findall(pattern, text):
            findings.append({"type": label, "value": match})

    account_patterns = [
        r"\b(?:account(?: number| no\.?)?|a/c(?: number| no\.?)?|bank account)\s*[:\-]?\s*(\d{9,18})\b",
        r"\b(\d{9,18})\s*(?:is|as)\s*(?:the )?(?:account(?: number| no\.?)|bank account)\b",
    ]
    for pattern in account_patterns:
        for match in re.findall(pattern, text, flags=re.IGNORECASE):
            findings.append({"type": "Bank Account", "value": match})
    return findings


def mask_sensitive_data(text):
    masked = text
    masks = {
        r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b": "[MASKED_EMAIL]",
        r"\b(?:\+91[- ]?)?[6-9]\d{9}\b": "[MASKED_PHONE]",
        r"\b[A-Z]{5}\d{4}[A-Z]\b": "[MASKED_PAN]",
        r"\b((?:account(?: number| no\.?)?|a/c(?: number| no\.?)?|bank account)\s*[:\-]?\s*)\d{9,18}\b": r"\1[MASKED_ACCOUNT]",
        r"\b\d{9,18}(\s*(?:is|as)\s*(?:the )?(?:account(?: number| no\.?)|bank account))\b": r"[MASKED_ACCOUNT]\1",
    }
    for pattern, replacement in masks.items():
        masked = re.sub(pattern, replacement, masked, flags=re.IGNORECASE)
    return masked


def detect_similar_clauses(text):
    sentences = split_sentences(text)
    similar_pairs = []
    for i in range(len(sentences)):
        for j in range(i + 1, len(sentences)):
            ratio = SequenceMatcher(None, sentences[i].lower(), sentences[j].lower()).ratio()
            if ratio >= 0.72 and sentences[i].lower() != sentences[j].lower():
                similar_pairs.append(
                    {
                        "sentence_a": sentences[i],
                        "sentence_b": sentences[j],
                        "similarity": round(ratio * 100, 2),
                    }
                )
    return similar_pairs[:8]


def summarize_contract(text, findings):
    sentences = split_sentences(text)
    if not sentences:
        return []
    clause_words = set()
    for patterns in CLAUSE_LIBRARY.values():
        clause_words.update(patterns)
    scored = []
    for sentence in sentences:
        lowered = sentence.lower()
        score = 0
        score += sum(1 for finding in findings if finding["term"] in lowered) * 3
        score += sum(1 for word in clause_words if contains_term(sentence, word))
        score += 2 if any(token in lowered for token in ["shall", "must", "liable", "terminate", "payment"]) else 0
        scored.append((score, sentence))
    top_sentences = [sentence for score, sentence in sorted(scored, key=lambda item: item[0], reverse=True)[:5] if score > 0]
    if not top_sentences:
        top_sentences = sentences[:3]
    return top_sentences


def run_compliance_checks(text):
    lowered = text.lower()
    issues = []
    for rule in COMPLIANCE_RULES:
        if any(trigger in lowered for trigger in rule["trigger"]) and not any(req in lowered for req in rule["required_any"]):
            issues.append(rule["message"])
    return issues


def build_recommendations(findings, ambiguity_findings, missing_clauses):
    recommendations = []
    seen = set()
    for finding in findings:
        suggestion = RISK_RECOMMENDATIONS.get(finding["term"])
        if suggestion and suggestion not in seen:
            recommendations.append({"title": finding["term"], "suggestion": suggestion})
            seen.add(suggestion)
    for item in ambiguity_findings[:3]:
        suggestion = f"Replace '{item['term']}' with measurable language such as a fixed timeline, threshold, or duty."
        if suggestion not in seen:
            recommendations.append({"title": item["term"], "suggestion": suggestion})
            seen.add(suggestion)
    for clause in missing_clauses:
        suggestion = f"Add a clearly drafted {clause.lower()} to reduce negotiation and enforcement gaps."
        if suggestion not in seen:
            recommendations.append({"title": clause, "suggestion": suggestion})
            seen.add(suggestion)
    return recommendations[:8]


def answer_question(question, text, entities, obligations, findings):
    question_lower = question.lower().strip()
    if not question_lower:
        return "Ask a question about payment, liability, termination, confidentiality, dates, or parties."
    if "who" in question_lower and obligations:
        first = obligations[0]
        return f"{first['subject']} is obligated to {first['action'].lower()} {first['object']}."
    if "date" in question_lower or "when" in question_lower:
        if entities["Dates"]:
            return f"Detected date references: {', '.join(entities['Dates'][:3])}."
    if "party" in question_lower or "parties" in question_lower:
        if entities["Parties"]:
            return f"Detected parties: {', '.join(entities['Parties'][:5])}."
    if "money" in question_lower or "amount" in question_lower or "payment" in question_lower:
        if entities["Money Values"]:
            return f"Detected money values: {', '.join(entities['Money Values'][:3])}."
    for topic, keywords in QUESTION_TOPICS.items():
        if topic in question_lower or any(keyword in question_lower for keyword in keywords):
            for sentence in split_sentences(text):
                if any(contains_term(sentence, keyword) for keyword in keywords):
                    return sentence
    if findings:
        strongest = sorted(findings, key=lambda item: level_priority(item["level"]))[0]
        return f"The strongest detected risk is '{strongest['term']}' under {strongest['level']}. {strongest['explanation']}"
    return "I could not find a direct answer in the contract text. Try asking about payment, liability, termination, or dates."


def compare_contracts(base_text, compare_text):
    base_findings = analyze_text(base_text)
    compare_findings = analyze_text(compare_text)
    base_terms = {item["term"] for item in base_findings}
    compare_terms = {item["term"] for item in compare_findings}
    base_summary = build_summary(base_findings, base_text, analyze_sentences(base_text))
    compare_summary = build_summary(compare_findings, compare_text, analyze_sentences(compare_text))
    if base_summary["overall_score"] < compare_summary["overall_score"]:
        safer_contract = "Base Contract"
        safety_reason = "The base contract has a lower overall risk score."
    elif compare_summary["overall_score"] < base_summary["overall_score"]:
        safer_contract = "Revised Contract"
        safety_reason = "The revised contract has a lower overall risk score."
    else:
        base_signal_count = len(base_findings) + len(base_summary["ambiguity"]) + len(base_summary["dependencies"]) + len(base_summary["compliance_issues"])
        compare_signal_count = len(compare_findings) + len(compare_summary["ambiguity"]) + len(compare_summary["dependencies"]) + len(compare_summary["compliance_issues"])
        if base_signal_count < compare_signal_count:
            safer_contract = "Base Contract"
            safety_reason = "Both scores are equal, but the base contract has fewer overall risk signals."
        elif compare_signal_count < base_signal_count:
            safer_contract = "Revised Contract"
            safety_reason = "Both scores are equal, but the revised contract has fewer overall risk signals."
        else:
            safer_contract = "Both are equally safe"
            safety_reason = "Both contracts have the same score and similar risk signal counts."
    clause_changes = []
    base_cards = {item["clause"]: item for item in base_summary["clause_cards"]}
    compare_cards = {item["clause"]: item for item in compare_summary["clause_cards"]}
    risk_rank = {"High Risk": 3, "Medium Risk": 2, "Low Risk": 1, "No Immediate Risk": 0, "Needs Review": 0}
    for clause_name in CLAUSE_LIBRARY:
        base_card = base_cards.get(clause_name, {"risk_level": "Needs Review", "score": 0})
        compare_card = compare_cards.get(clause_name, {"risk_level": "Needs Review", "score": 0})
        delta = compare_card["score"] - base_card["score"]
        if delta < 0:
            change = "Improved"
            summary = f"{clause_name} looks safer in the revised contract."
        elif delta > 0:
            change = "Worsened"
            summary = f"{clause_name} appears riskier in the revised contract."
        else:
            base_rank = risk_rank.get(base_card["risk_level"], 0)
            compare_rank = risk_rank.get(compare_card["risk_level"], 0)
            if compare_rank < base_rank:
                change = "Improved"
                summary = f"{clause_name} risk level improved in the revised contract."
            elif compare_rank > base_rank:
                change = "Worsened"
                summary = f"{clause_name} risk level worsened in the revised contract."
            else:
                continue
        clause_changes.append(
            {
                "clause": clause_name,
                "change": change,
                "summary": summary,
                "base_risk": base_card["risk_level"],
                "revised_risk": compare_card["risk_level"],
            }
        )
    return {
        "similarity": round(SequenceMatcher(None, normalize_text(base_text).lower(), normalize_text(compare_text).lower()).ratio() * 100, 2),
        "base_score": base_summary["overall_score"],
        "compare_score": compare_summary["overall_score"],
        "base_risk": base_summary["overall_risk"],
        "compare_risk": compare_summary["overall_risk"],
        "safer_contract": safer_contract,
        "safety_reason": safety_reason,
        "added_risks": sorted(compare_terms - base_terms),
        "removed_risks": sorted(base_terms - compare_terms),
        "clause_changes": clause_changes,
    }


def build_summary(findings, text, sentence_findings):
    counts = {"High Risk": 0, "Medium Risk": 0, "Low Risk": 0}
    weights = {"High Risk": 0, "Medium Risk": 0, "Low Risk": 0}
    for finding in findings:
        counts[finding["level"]] += 1
        weights[finding["level"]] += finding["weight"]

    ambiguity_findings = detect_ambiguity(text)
    obligations = extract_obligations(text)
    dependencies = detect_clause_dependencies(text)
    clauses = detect_clause_types(text)
    clause_cards = build_clause_cards(clauses, findings, text)
    missing_clauses = [item for item in MISSING_CLAUSE_TARGETS if item not in {clause["clause"] for clause in clauses}]
    entity_map = extract_entities(text)
    sensitive_data = detect_sensitive_data(text)
    similar_clauses = detect_similar_clauses(text)
    compliance_issues = run_compliance_checks(text)

    has_material_risk_signal = any(
        [
            findings,
            ambiguity_findings,
            compliance_issues,
            dependencies,
        ]
    )

    keyword_score = weights["High Risk"] + weights["Medium Risk"] + weights["Low Risk"]
    ambiguity_score = min(len(ambiguity_findings) * 4, 12)
    compliance_score = min(len(compliance_issues) * 8, 16)
    dependency_score = min(len(dependencies) * 6, 12)
    sentence_score = min(len(sentence_findings) * 4, 12)
    raw_score = keyword_score + ambiguity_score + compliance_score + dependency_score
    if has_material_risk_signal:
        overall_score = min(raw_score + sentence_score, 100)
    else:
        overall_score = 0

    if overall_score >= 80:
        overall_risk = "High Risk"
        overall_class = "high"
    elif overall_score >= 50:
        overall_risk = "Medium Risk"
        overall_class = "medium"
    elif overall_score > 0:
        overall_risk = "Low Risk"
        overall_class = "low"
    else:
        overall_risk = "No Immediate Risk"
        overall_class = "neutral"

    top_terms = Counter(item["term"] for item in findings).most_common(5)
    recommendations = build_recommendations(findings, ambiguity_findings, missing_clauses)
    return {
        "counts": counts,
        "weights": weights,
        "overall_risk": overall_risk,
        "overall_class": overall_class,
        "overall_score": overall_score,
        "total_matches": sum(counts.values()),
        "total_sentences": len(split_sentences(text)),
        "risky_sentences": len(sentence_findings),
        "clauses": clauses,
        "clause_cards": clause_cards,
        "missing_clauses": missing_clauses,
        "entities": entity_map,
        "top_terms": top_terms,
        "ambiguity": ambiguity_findings,
        "obligations": obligations,
        "dependencies": dependencies,
        "sensitive_data": sensitive_data,
        "similar_clauses": similar_clauses,
        "compliance_issues": compliance_issues,
        "recommendations": recommendations,
        "summary_points": summarize_contract(text, findings),
        "score_breakdown": {
            "keywords": keyword_score,
            "ambiguity": ambiguity_score,
            "compliance": compliance_score,
            "dependency": dependency_score,
            "sentence_context": sentence_score,
        },
        "analysis_note": "This is a rule-based legal screening result built from keywords, ambiguity checks, clause patterns, and simple dependency logic.",
    }


def explain_contract(summary, findings):
    if not findings and not summary["ambiguity"] and not summary["compliance_issues"]:
        return "No tracked legal risk indicators were found in the submitted contract text."
    reasons = [item["explanation"] for item in sorted(findings, key=lambda item: level_priority(item["level"]))[:3]]
    if summary["ambiguity"]:
        reasons.append(f"Ambiguous wording like '{summary['ambiguity'][0]['term']}' reduces clarity.")
    if summary["compliance_issues"]:
        reasons.append(summary["compliance_issues"][0])
    joined = " ".join(reasons[:4])
    return f"This contract is marked as {summary['overall_risk']} with a score of {summary['overall_score']}/100. {joined}"


def _apply_term_highlights(rendered_text):
    highlighted = rendered_text
    replacements = []
    for level, config in RISK_LIBRARY.items():
        for term in sorted(config["keywords"], key=len, reverse=True):
            replacements.append((term, level_class(level), level))
    for term, css_class, label in replacements:
        pattern = re.compile(rf"\b({re.escape(term)})\b", re.IGNORECASE)
        highlighted = pattern.sub(
            lambda match: f"<span class='highlight {css_class}' title='{html.escape(label)}'>{html.escape(match.group(0))}</span>",
            highlighted,
        )
    for term in sorted(AMBIGUOUS_TERMS, key=len, reverse=True):
        pattern = re.compile(rf"\b({re.escape(term)})\b", re.IGNORECASE)
        highlighted = pattern.sub(
            lambda match: f"<span class='highlight ambiguous' title='Ambiguous'>{html.escape(match.group(0))}</span>",
            highlighted,
        )
    return highlighted


def highlight_risk_terms(text):
    highlighted = _apply_term_highlights(html.escape(text))
    return highlighted.replace("\n", "<br>")


def highlight_sentences_with_tooltips(text, sentence_findings):
    rendered = html.escape(text)
    ordered_findings = sorted(sentence_findings, key=lambda item: len(item["sentence"]), reverse=True)
    for item in ordered_findings:
        reason = html.escape(item["reason"] or f"{item['level']} signal detected")
        sentence_html = html.escape(item["sentence"])
        rendered = rendered.replace(
            sentence_html,
            f"<span class='sentence-highlight {level_class(item['level'])}' title='{reason}'>{sentence_html}</span>",
            1,
        )
    rendered = _apply_term_highlights(rendered)
    return rendered.replace("\n", "<br>")


def build_redline_view(text):
    rendered = html.escape(text)
    for risky_term, suggestion in RISK_RECOMMENDATIONS.items():
        safe_phrase = html.escape(suggestion.split(".")[0])
        pattern = re.compile(rf"\b({re.escape(risky_term)})\b", re.IGNORECASE)
        rendered = pattern.sub(
            lambda match: f"<span class='strike'>{html.escape(match.group(0))}</span> <span class='insert'>{safe_phrase}</span>",
            rendered,
        )
    return rendered.replace("\n", "<br>")


def role_based_summary(role, summary, findings):
    if role == "Manager":
        top_issue = findings[0]["term"] if findings else "no major tracked issue"
        return f"Risk score {summary['overall_score']}/100. Overall status: {summary['overall_risk']}. Main concern: {top_issue}."
    if role == "Client":
        missing = ", ".join(summary["missing_clauses"]) if summary["missing_clauses"] else "no major missing clauses"
        return f"This contract may affect your obligations and costs. It is rated {summary['overall_risk']} and appears to have {missing}."
    return explain_contract(summary, findings)


def build_report_text(text, findings, sentence_findings, summary):
    lines = [
        "Legal Contract Risk Analyzer Report",
        "",
        f"Overall Risk: {summary['overall_risk']}",
        f"Overall Risk Score: {summary['overall_score']}/100",
        f"Total Sentences: {summary['total_sentences']}",
        f"Risky Sentences: {summary['risky_sentences']}",
        f"Total Risk Matches: {summary['total_matches']}",
        "",
        "Key Summary Points:",
    ]
    lines.extend(f"- {point}" for point in summary["summary_points"] or ["No summary points available"])
    lines.extend(["", "Clause Dependencies:"])
    if summary["dependencies"]:
        lines.extend(f"- {item['source']} depends on {item['depends_on']}: {item['sentence']}" for item in summary["dependencies"])
    else:
        lines.append("- No clause dependencies detected")
    lines.extend(["", "Ambiguity Findings:"])
    if summary["ambiguity"]:
        lines.extend(f"- {item['term']}: {item['message']}" for item in summary["ambiguity"])
    else:
        lines.append("- No major ambiguous terms detected")
    lines.extend(["", "Obligations:"])
    if summary["obligations"]:
        lines.extend(f"- {item['subject']} must {item['action'].lower()} {item['object']}" for item in summary["obligations"])
    else:
        lines.append("- No obligations extracted")
    lines.extend(["", "Compliance Issues:"])
    if summary["compliance_issues"]:
        lines.extend(f"- {issue}" for issue in summary["compliance_issues"])
    else:
        lines.append("- No compliance gaps detected by the rule set")
    lines.extend(["", "Sensitive Data:"])
    if summary["sensitive_data"]:
        lines.extend(f"- {item['type']}: {item['value']}" for item in summary["sensitive_data"])
    else:
        lines.append("- No sensitive data patterns detected")
    lines.extend(["", "Recommendations:"])
    if summary["recommendations"]:
        lines.extend(f"- {item['title']}: {item['suggestion']}" for item in summary["recommendations"])
    else:
        lines.append("- No recommendations generated")
    lines.extend(["", "Explainable Summary:", explain_contract(summary, findings), "", "Original Contract Text:", text or "No text submitted"])
    return "\n".join(lines)


def build_pdf_report(report_text):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=0.6 * inch, leftMargin=0.6 * inch, topMargin=0.6 * inch, bottomMargin=0.6 * inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleStyle", parent=styles["Heading1"], textColor=HexColor("#1f2937"), fontSize=18, leading=22, spaceAfter=12)
    body_style = ParagraphStyle("BodyStyle", parent=styles["BodyText"], fontSize=10, leading=14, textColor=HexColor("#334155"))
    story = []
    lines = report_text.split("\n")
    if lines:
        story.append(Paragraph(html.escape(lines[0]), title_style))
        story.append(Spacer(1, 0.15 * inch))
    for line in lines[1:]:
        story.append(Paragraph(html.escape(line) if line.strip() else "&nbsp;", body_style))
        story.append(Spacer(1, 0.06 * inch))
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def build_docx_report(report_text):
    document = Document()
    lines = report_text.split("\n")
    if lines:
        document.add_heading(lines[0], level=0)
    for line in lines[1:]:
        if line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            document.add_paragraph(line)
        else:
            document.add_paragraph("")
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def run_full_analysis(text, source_name="Pasted Contract"):
    findings = analyze_text(text) if text else []
    sentence_findings = analyze_sentences(text) if text else []
    summary = build_summary(findings, text, sentence_findings)
    report_text = build_report_text(text, findings, sentence_findings, summary)
    pdf_report = build_pdf_report(report_text) if text else b""
    docx_report = build_docx_report(report_text) if text else b""
    return {
        "source_name": source_name,
        "text": text,
        "findings": findings,
        "sentence_findings": sentence_findings,
        "summary": summary,
        "report_text": report_text,
        "pdf_report": pdf_report,
        "docx_report": docx_report,
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
